"""Document processing module for handling various file formats."""

import os
import tempfile
from pathlib import Path
from typing import List, Optional
import logging

from pypdf import PdfReader
from docx import Document as DocxDocument
import markdown

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Processes documents in various formats (PDF, Word, Markdown)."""

    SUPPORTED_FORMATS = {".pdf", ".docx", ".md", ".txt"}
    MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

    @staticmethod
    def validate_file(file_path: str) -> bool:
        """Validate if file exists and is in supported format."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if path.suffix.lower() not in DocumentProcessor.SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported file format: {path.suffix}. "
                f"Supported formats: {DocumentProcessor.SUPPORTED_FORMATS}"
            )

        file_size = path.stat().st_size
        if file_size > DocumentProcessor.MAX_FILE_SIZE:
            raise ValueError(
                f"File size {file_size} exceeds maximum allowed size "
                f"{DocumentProcessor.MAX_FILE_SIZE}"
            )

        return True

    @staticmethod
    def process_document(file_path: str) -> str:
        """
        Process a document and extract text content.

        Args:
            file_path: Path to the document file

        Returns:
            Extracted text content

        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported or file is too large
            Exception: If document processing fails
        """
        DocumentProcessor.validate_file(file_path)

        path = Path(file_path)
        suffix = path.suffix.lower()

        try:
            if suffix == ".pdf":
                return DocumentProcessor._parse_pdf(file_path)
            elif suffix == ".docx":
                return DocumentProcessor._parse_docx(file_path)
            elif suffix == ".md":
                return DocumentProcessor._parse_markdown(file_path)
            elif suffix == ".txt":
                return DocumentProcessor._parse_text(file_path)
            else:
                raise ValueError(f"Unsupported file format: {suffix}")
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise

    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """Parse PDF file and extract text."""
        try:
            reader = PdfReader(file_path)
            text_parts = []

            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(
                        f"Failed to extract text from page {page_num} in {file_path}: {e}"
                    )

            if not text_parts:
                raise ValueError("No text could be extracted from PDF")

            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            raise

    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """Parse Word document and extract text."""
        try:
            doc = DocxDocument(file_path)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            if not text_parts:
                raise ValueError("No text could be extracted from Word document")

            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error parsing Word document {file_path}: {str(e)}")
            raise

    @staticmethod
    def _parse_markdown(file_path: str) -> str:
        """Parse Markdown file and extract text."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            if not content.strip():
                raise ValueError("Markdown file is empty")

            # Convert markdown to HTML then extract text
            # For now, just return the raw markdown content
            return content
        except Exception as e:
            logger.error(f"Error parsing Markdown file {file_path}: {str(e)}")
            raise

    @staticmethod
    def _parse_text(file_path: str) -> str:
        """Parse plain text file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            if not content.strip():
                raise ValueError("Text file is empty")

            return content
        except Exception as e:
            logger.error(f"Error parsing text file {file_path}: {str(e)}")
            raise
