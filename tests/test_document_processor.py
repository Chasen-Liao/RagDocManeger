"""Tests for document processor."""

import pytest
import tempfile
from pathlib import Path
from RagDocMan.rag.document_processor import DocumentProcessor


class TestDocumentProcessorValidation:
    """Test document validation."""

    def test_validate_file_not_found(self):
        """Test validation fails for non-existent file."""
        with pytest.raises(FileNotFoundError):
            DocumentProcessor.validate_file("/nonexistent/file.pdf")

    def test_validate_unsupported_format(self):
        """Test validation fails for unsupported format."""
        with tempfile.NamedTemporaryFile(suffix=".xyz") as f:
            with pytest.raises(ValueError, match="Unsupported file format"):
                DocumentProcessor.validate_file(f.name)

    def test_validate_file_too_large(self):
        """Test validation fails for files exceeding size limit."""
        with tempfile.NamedTemporaryFile(suffix=".txt") as f:
            # Write data larger than max size
            large_data = b"x" * (DocumentProcessor.MAX_FILE_SIZE + 1)
            f.write(large_data)
            f.flush()

            with pytest.raises(ValueError, match="exceeds maximum allowed size"):
                DocumentProcessor.validate_file(f.name)

    def test_validate_supported_formats(self):
        """Test validation passes for supported formats."""
        for ext in [".pdf", ".docx", ".md", ".txt"]:
            with tempfile.NamedTemporaryFile(suffix=ext) as f:
                f.write(b"test content")
                f.flush()
                assert DocumentProcessor.validate_file(f.name) is True


class TestTextFileProcessing:
    """Test plain text file processing."""

    def test_parse_text_file(self):
        """Test parsing plain text file."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
            f.write("Hello World\nThis is a test")
            f.flush()
            temp_path = f.name

        try:
            content = DocumentProcessor.process_document(temp_path)
            assert "Hello World" in content
            assert "This is a test" in content
        finally:
            Path(temp_path).unlink()

    def test_parse_empty_text_file(self):
        """Test parsing empty text file raises error."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
            f.write("")
            f.flush()
            temp_path = f.name

        try:
            with pytest.raises(ValueError, match="Text file is empty"):
                DocumentProcessor.process_document(temp_path)
        finally:
            Path(temp_path).unlink()

    def test_parse_text_with_unicode(self):
        """Test parsing text file with unicode characters."""
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False, encoding="utf-8"
        ) as f:
            f.write("Hello 世界 🌍")
            f.flush()
            temp_path = f.name

        try:
            content = DocumentProcessor.process_document(temp_path)
            assert "Hello" in content
            assert "世界" in content
        finally:
            Path(temp_path).unlink()


class TestMarkdownProcessing:
    """Test Markdown file processing."""

    def test_parse_markdown_file(self):
        """Test parsing Markdown file."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
            f.write("# Title\n\nThis is a paragraph.\n\n## Section\n\nMore content.")
            f.flush()
            temp_path = f.name

        try:
            content = DocumentProcessor.process_document(temp_path)
            assert "Title" in content
            assert "This is a paragraph" in content
            assert "Section" in content
        finally:
            Path(temp_path).unlink()

    def test_parse_empty_markdown_file(self):
        """Test parsing empty Markdown file raises error."""
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
            f.write("")
            f.flush()
            temp_path = f.name

        try:
            with pytest.raises(ValueError, match="Markdown file is empty"):
                DocumentProcessor.process_document(temp_path)
        finally:
            Path(temp_path).unlink()


class TestWordDocumentProcessing:
    """Test Word document processing."""

    def test_parse_docx_file(self):
        """Test parsing Word document."""
        from docx import Document as DocxDocument

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        try:
            # Create a test Word document
            doc = DocxDocument()
            doc.add_paragraph("Hello World")
            doc.add_paragraph("This is a test document")
            doc.save(temp_path)

            content = DocumentProcessor.process_document(temp_path)
            assert "Hello World" in content
            assert "This is a test document" in content
        finally:
            Path(temp_path).unlink()

    def test_parse_docx_with_table(self):
        """Test parsing Word document with table."""
        from docx import Document as DocxDocument

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            temp_path = f.name

        try:
            doc = DocxDocument()
            doc.add_paragraph("Document with table")
            table = doc.add_table(rows=2, cols=2)
            table.rows[0].cells[0].text = "Header 1"
            table.rows[0].cells[1].text = "Header 2"
            table.rows[1].cells[0].text = "Data 1"
            table.rows[1].cells[1].text = "Data 2"
            doc.save(temp_path)

            content = DocumentProcessor.process_document(temp_path)
            assert "Document with table" in content
            assert "Header 1" in content
            assert "Data 1" in content
        finally:
            Path(temp_path).unlink()


class TestPdfProcessing:
    """Test PDF file processing."""

    def test_parse_pdf_file(self):
        """Test parsing PDF file."""
        from pypdf import PdfWriter
        from io import BytesIO

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            temp_path = f.name

        try:
            # Create a simple PDF
            writer = PdfWriter()
            writer.add_blank_page(width=200, height=200)
            with open(temp_path, "wb") as output_file:
                writer.write(output_file)

            # Note: This PDF won't have extractable text, but we test the parsing works
            content = DocumentProcessor.process_document(temp_path)
            assert isinstance(content, str)
        except ValueError as e:
            # It's OK if no text can be extracted from blank PDF
            assert "No text could be extracted" in str(e)
        finally:
            Path(temp_path).unlink()


class TestDocumentProcessorIntegration:
    """Integration tests for document processor."""

    def test_process_document_with_invalid_path(self):
        """Test processing with invalid path."""
        with pytest.raises(FileNotFoundError):
            DocumentProcessor.process_document("/invalid/path/file.txt")

    def test_process_document_with_unsupported_format(self):
        """Test processing with unsupported format."""
        with tempfile.NamedTemporaryFile(suffix=".xyz") as f:
            with pytest.raises(ValueError, match="Unsupported file format"):
                DocumentProcessor.process_document(f.name)

    def test_supported_formats_constant(self):
        """Test supported formats are defined."""
        assert ".pdf" in DocumentProcessor.SUPPORTED_FORMATS
        assert ".docx" in DocumentProcessor.SUPPORTED_FORMATS
        assert ".md" in DocumentProcessor.SUPPORTED_FORMATS
        assert ".txt" in DocumentProcessor.SUPPORTED_FORMATS
